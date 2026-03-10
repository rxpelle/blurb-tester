"""Tests for the docx parser."""

import os
import pytest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from book_formatter.parsers.docx_parser import (
    docx_to_markdown,
    _heading_level,
    _is_scene_break,
    _runs_to_markdown,
)
from book_formatter.parsers.markdown import parse_manuscript
from book_formatter.config import BookConfig


def _create_docx(filepath, paragraphs):
    """Helper to create a .docx file with structured content.

    paragraphs: list of dicts with keys:
      - text: str
      - style: str (optional, default 'Normal')
      - bold: bool (optional)
      - italic: bool (optional)
      - alignment: WD_ALIGN_PARAGRAPH (optional)
    """
    doc = Document()
    for p in paragraphs:
        if isinstance(p, str):
            p = {'text': p}
        para = doc.add_paragraph()
        if 'style' in p:
            para.style = p['style']
        if 'alignment' in p:
            para.alignment = p['alignment']
        run = para.add_run(p.get('text', ''))
        if p.get('bold'):
            run.bold = True
        if p.get('italic'):
            run.italic = True
    doc.save(filepath)


class TestDocxToMarkdown:
    """Tests for docx to markdown conversion."""

    def test_basic_paragraphs(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        _create_docx(filepath, [
            'First paragraph.',
            'Second paragraph.',
        ])
        md = docx_to_markdown(filepath)
        assert 'First paragraph.' in md
        assert 'Second paragraph.' in md

    def test_headings_converted(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading('Chapter One', level=1)
        doc.add_paragraph('Content under chapter one.')
        doc.add_heading('Section Two', level=2)
        doc.add_paragraph('Content under section.')
        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '# Chapter One' in md
        assert '## Section Two' in md

    def test_bold_formatting(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        para = doc.add_paragraph()
        para.add_run('Normal text and ')
        bold_run = para.add_run('bold text')
        bold_run.bold = True
        para.add_run(' here.')
        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '**bold text**' in md

    def test_italic_formatting(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        para = doc.add_paragraph()
        para.add_run('She said ')
        italic_run = para.add_run('something important')
        italic_run.italic = True
        para.add_run('.')
        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '*something important*' in md

    def test_bold_italic_formatting(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('emphasized')
        run.bold = True
        run.italic = True
        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '***emphasized***' in md

    def test_scene_break_detected(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_paragraph('Before the break.')
        para = doc.add_paragraph('* * *')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('After the break.')
        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '---' in md

    def test_full_manuscript_conversion(self, tmp_dir):
        """Test a realistic multi-chapter manuscript."""
        filepath = os.path.join(tmp_dir, 'manuscript.docx')
        doc = Document()

        doc.add_heading('The Storm', level=1)
        doc.add_paragraph('Lightning cracked the sky. Rain poured down in sheets.')
        doc.add_paragraph('Maria clutched her bag and ran for cover.')
        para = doc.add_paragraph('* * *')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('When the storm passed, nothing remained.')

        doc.add_heading('The Aftermath', level=1)
        doc.add_paragraph('The town lay in ruins. Survivors emerged slowly.')

        doc.add_heading('Epilogue', level=1)
        doc.add_paragraph('Years later, they rebuilt.')

        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '# The Storm' in md
        assert '# The Aftermath' in md
        assert '# Epilogue' in md
        assert '---' in md
        assert 'Lightning cracked' in md

    def test_excessive_blank_lines_collapsed(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_paragraph('First.')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('Second.')
        doc.save(filepath)

        md = docx_to_markdown(filepath)
        assert '\n\n\n\n' not in md


class TestDocxParseManuscript:
    """Integration tests for .docx files through parse_manuscript."""

    def test_docx_parsed_as_book(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'manuscript.docx')
        doc = Document()
        doc.add_heading('Chapter One', level=1)
        doc.add_paragraph('Content of chapter one. ' * 20)
        doc.add_heading('Chapter Two', level=1)
        doc.add_paragraph('Content of chapter two. ' * 15)
        doc.save(filepath)

        config = BookConfig()
        config._config_dir = tmp_dir
        config.manuscript = 'manuscript.docx'

        book = parse_manuscript(config)
        assert len(book.chapters) == 2
        assert book.chapters[0].title == 'Chapter One'
        assert book.chapters[1].title == 'Chapter Two'

    def test_docx_source_file_points_to_original(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading('Test', level=1)
        doc.add_paragraph('Content.')
        doc.save(filepath)

        config = BookConfig()
        config._config_dir = tmp_dir
        config.manuscript = 'test.docx'

        book = parse_manuscript(config)
        assert book.chapters[0].source_file.endswith('.docx')

    def test_docx_special_chapters_detected(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading('Prologue', level=1)
        doc.add_paragraph('Before the story.')
        doc.add_heading('The Story', level=1)
        doc.add_paragraph('Main content.')
        doc.add_heading('Epilogue', level=1)
        doc.add_paragraph('After the story.')
        doc.save(filepath)

        config = BookConfig()
        config._config_dir = tmp_dir
        config.manuscript = 'test.docx'

        book = parse_manuscript(config)
        assert book.chapters[0].is_unnumbered is True  # Prologue
        assert book.chapters[1].is_unnumbered is False  # Normal
        assert book.chapters[2].is_epilogue is True  # Epilogue

    def test_docx_word_counts(self, tmp_dir):
        filepath = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading('Chapter One', level=1)
        doc.add_paragraph('Word ' * 100)
        doc.save(filepath)

        config = BookConfig()
        config._config_dir = tmp_dir
        config.manuscript = 'test.docx'

        book = parse_manuscript(config)
        assert book.chapters[0].word_count > 50
        assert book.word_count > 50


class TestHeadingLevel:
    """Tests for heading level extraction."""

    def test_heading_1(self):
        assert _heading_level('Heading 1') == 1

    def test_heading_3(self):
        assert _heading_level('Heading 3') == 3

    def test_heading_capped_at_6(self):
        assert _heading_level('Heading 9') == 6

    def test_no_number_defaults_to_1(self):
        assert _heading_level('Heading') == 1


class TestIsSceneBreak:
    """Tests for scene break detection."""

    def test_asterisk_break(self, tmp_dir):
        doc = Document()
        para = doc.add_paragraph('* * *')
        assert _is_scene_break(para) is True

    def test_triple_asterisk(self, tmp_dir):
        doc = Document()
        para = doc.add_paragraph('***')
        assert _is_scene_break(para) is True

    def test_dashes(self, tmp_dir):
        doc = Document()
        para = doc.add_paragraph('---')
        assert _is_scene_break(para) is True

    def test_normal_text_not_break(self, tmp_dir):
        doc = Document()
        para = doc.add_paragraph('This is a normal paragraph.')
        assert _is_scene_break(para) is False

    def test_empty_not_break(self, tmp_dir):
        doc = Document()
        para = doc.add_paragraph('')
        assert _is_scene_break(para) is False


class TestRunsToMarkdown:
    """Tests for inline formatting conversion."""

    def test_plain_text(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run('Hello world')
        assert _runs_to_markdown(para.runs) == 'Hello world'

    def test_bold(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('bold')
        run.bold = True
        assert _runs_to_markdown(para.runs) == '**bold**'

    def test_italic(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('italic')
        run.italic = True
        assert _runs_to_markdown(para.runs) == '*italic*'

    def test_bold_italic(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('both')
        run.bold = True
        run.italic = True
        assert _runs_to_markdown(para.runs) == '***both***'

    def test_mixed_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run('normal ')
        bold = para.add_run('bold')
        bold.bold = True
        para.add_run(' normal')
        result = _runs_to_markdown(para.runs)
        assert 'normal ' in result
        assert '**bold**' in result

    def test_empty_runs(self):
        assert _runs_to_markdown([]) == ''
